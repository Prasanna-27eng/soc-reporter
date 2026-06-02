import json
import io
from datetime import datetime
from typing import Any, Dict, List
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, PageBreak, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Color palette ────────────────────────────────────────────────────────────
NAVY      = colors.HexColor("#0D1B2A")
CYAN      = colors.HexColor("#00B4D8")
SLATE     = colors.HexColor("#1E2D3D")
LIGHT_BG  = colors.HexColor("#F0F4F8")
WHITE     = colors.white
RED_CRIT  = colors.HexColor("#E63946")
ORANGE_HI = colors.HexColor("#F4A261")
YELLOW_MD = colors.HexColor("#E9C46A")
GREEN_LO  = colors.HexColor("#2A9D8F")
GRAY_TEXT = colors.HexColor("#4A5568")
MID_GRAY  = colors.HexColor("#CBD5E0")

SEVERITY_COLORS = {
    "Critical": RED_CRIT,
    "High": ORANGE_HI,
    "Medium": YELLOW_MD,
    "Low": GREEN_LO,
}

VERDICT_COLORS = {
    "malicious": RED_CRIT,
    "suspicious": ORANGE_HI,
    "clean": GREEN_LO,
    "not_found": GRAY_TEXT,
    "unknown": MID_GRAY,
    "error": MID_GRAY,
}

def severity_color(s):
    return SEVERITY_COLORS.get(s, GRAY_TEXT)

def verdict_color(v):
    return VERDICT_COLORS.get(v, MID_GRAY)

# ─── PDF helpers ──────────────────────────────────────────────────────────────
def build_styles():
    styles = getSampleStyleSheet()
    base = {"fontName": "Helvetica", "textColor": NAVY}

    custom = {
        "CoverTitle": ParagraphStyle("CoverTitle", fontSize=28, leading=34, textColor=WHITE, fontName="Helvetica-Bold", alignment=TA_LEFT),
        "CoverSub": ParagraphStyle("CoverSub", fontSize=13, leading=18, textColor=colors.HexColor("#90CAF9"), fontName="Helvetica", alignment=TA_LEFT),
        "CoverMeta": ParagraphStyle("CoverMeta", fontSize=10, leading=14, textColor=colors.HexColor("#CFD8DC"), fontName="Helvetica", alignment=TA_LEFT),
        "SectionHeader": ParagraphStyle("SectionHeader", fontSize=13, leading=18, textColor=WHITE, fontName="Helvetica-Bold", alignment=TA_LEFT, spaceBefore=4, spaceAfter=4),
        "BodyText": ParagraphStyle("BodyText", fontSize=10, leading=15, textColor=NAVY, fontName="Helvetica", alignment=TA_JUSTIFY, spaceBefore=4, spaceAfter=4),
        "SmallLabel": ParagraphStyle("SmallLabel", fontSize=8, leading=11, textColor=GRAY_TEXT, fontName="Helvetica"),
        "TableHeader": ParagraphStyle("TableHeader", fontSize=9, leading=12, textColor=WHITE, fontName="Helvetica-Bold"),
        "TableCell": ParagraphStyle("TableCell", fontSize=9, leading=12, textColor=NAVY, fontName="Helvetica"),
        "MitreTag": ParagraphStyle("MitreTag", fontSize=8, leading=11, textColor=WHITE, fontName="Helvetica-Bold"),
        "Mono": ParagraphStyle("Mono", fontSize=8, leading=12, textColor=colors.HexColor("#1A1A2E"), fontName="Courier", backColor=colors.HexColor("#EEF2F7"), leftIndent=4, rightIndent=4),
    }
    return {**{k: styles[k] for k in styles.byName}, **custom}

def section_banner(title, styles):
    data = [[Paragraph(f"  {title}", styles["SectionHeader"])]]
    t = Table(data, colWidths=[170*mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), SLATE),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [SLATE]),
    ]))
    return t

def severity_badge(severity, styles):
    col = severity_color(severity)
    data = [[Paragraph(f"  {severity.upper()}  ", ParagraphStyle("Badge", fontSize=11, textColor=WHITE, fontName="Helvetica-Bold"))]]
    t = Table(data, colWidths=[40*mm])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), col), ("ROUNDEDCORNERS", [4, 4, 4, 4])]))
    return t

def ioc_table(ioc_list, vt_results, styles):
    header = [
        Paragraph("IOC", styles["TableHeader"]),
        Paragraph("Type", styles["TableHeader"]),
        Paragraph("VT Verdict", styles["TableHeader"]),
        Paragraph("Detections", styles["TableHeader"]),
    ]
    rows = [header]
    for ioc in ioc_list:
        vt = vt_results.get(ioc, {})
        verdict = vt.get("verdict", "not_checked")
        det = f"{vt.get('malicious_count', '-')}/{vt.get('total_engines', '-')}" if vt.get("total_engines") else "—"
        vc = verdict_color(verdict)
        rows.append([
            Paragraph(ioc[:45], styles["TableCell"]),
            Paragraph(vt.get("type", "—"), styles["TableCell"]),
            Paragraph(verdict.upper(), ParagraphStyle("VCell", fontSize=9, textColor=WHITE, fontName="Helvetica-Bold")),
            Paragraph(det, styles["TableCell"]),
        ])

    col_widths = [75*mm, 22*mm, 30*mm, 25*mm]
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    style = [
        ("BACKGROUND", (0, 0), (-1, 0), SLATE),
        ("BACKGROUND", (0, 1), (-1, -1), LIGHT_BG),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT_BG]),
        ("GRID", (0, 0), (-1, -1), 0.3, MID_GRAY),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]
    for i, ioc in enumerate(ioc_list):
        vt = vt_results.get(ioc, {})
        verdict = vt.get("verdict", "not_checked")
        vc = verdict_color(verdict)
        style.append(("BACKGROUND", (2, i + 1), (2, i + 1), vc))
    t.setStyle(TableStyle(style))
    return t

def mitre_table(techniques, styles):
    if not techniques:
        return Paragraph("No MITRE ATT&CK techniques mapped.", styles["BodyText"])
    header = [Paragraph(h, styles["TableHeader"]) for h in ["ID", "Technique", "Tactic", "Evidence"]]
    rows = [header]
    for t in techniques:
        rows.append([
            Paragraph(t.get("technique_id", ""), styles["TableCell"]),
            Paragraph(t.get("technique_name", ""), styles["TableCell"]),
            Paragraph(t.get("tactic", ""), styles["TableCell"]),
            Paragraph((t.get("evidence", ""))[:80], styles["TableCell"]),
        ])
    col_widths = [20*mm, 50*mm, 35*mm, 65*mm]
    tbl = Table(rows, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), SLATE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT_BG]),
        ("GRID", (0, 0), (-1, -1), 0.3, MID_GRAY),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#E8F4FD")),
    ]))
    return tbl

def timeline_table(events, styles):
    if not events:
        return Paragraph("No timeline events recorded.", styles["BodyText"])
    rows = [[Paragraph("Event", styles["TableHeader"])]]
    for ev in events:
        rows.append([Paragraph(str(ev), styles["TableCell"])])
    tbl = Table(rows, colWidths=[170*mm], repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), SLATE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT_BG]),
        ("GRID", (0, 0), (-1, -1), 0.3, MID_GRAY),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 1), (0, -1), 12),
    ]))
    return tbl

def header_footer(canvas, doc, case):
    canvas.saveState()
    w, h = A4
    # Header bar
    canvas.setFillColor(NAVY)
    canvas.rect(0, h - 18*mm, w, 18*mm, fill=1, stroke=0)
    canvas.setFillColor(CYAN)
    canvas.rect(0, h - 19*mm, w, 1.5*mm, fill=1, stroke=0)
    canvas.setFont("Helvetica-Bold", 9)
    canvas.setFillColor(WHITE)
    canvas.drawString(15*mm, h - 12*mm, f"SOC Incident Report  |  {case.get('case_number', '')}  |  {case.get('classification', 'TLP:AMBER')}")
    canvas.setFont("Helvetica", 8)
    canvas.drawRightString(w - 15*mm, h - 12*mm, f"CONFIDENTIAL  —  {datetime.utcnow().strftime('%Y-%m-%d')}")
    # Footer
    canvas.setFillColor(NAVY)
    canvas.rect(0, 0, w, 10*mm, fill=1, stroke=0)
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#90CAF9"))
    canvas.drawString(15*mm, 3.5*mm, f"Prepared by: {case.get('analyst_name', 'SOC Analyst')}  |  Customer: {case.get('customer_name', 'N/A')}")
    canvas.drawRightString(w - 15*mm, 3.5*mm, f"Page {doc.page}")
    canvas.restoreState()

# ─── PDF Generator ────────────────────────────────────────────────────────────
def _parse_json_field(val, default):
    if isinstance(val, (list, dict)):
        return val
    try:
        return json.loads(val or ('"[]"' if isinstance(default, list) else '{}'))
    except Exception:
        return default

def generate_pdf(case: Dict) -> bytes:
    buf = io.BytesIO()
    styles = build_styles()
    iocs = _parse_json_field(case.get("iocs", "[]"), [])
    timeline = _parse_json_field(case.get("timeline_events", "[]"), [])
    mitre = _parse_json_field(case.get("mitre_techniques", "[]"), [])
    vt_results = _parse_json_field(case.get("vt_results", "{}"), {})

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=15*mm, rightMargin=15*mm,
        topMargin=25*mm, bottomMargin=18*mm,
        title=f"SOC Report - {case.get('case_number', '')}",
    )
    doc.onFirstPage = lambda c, d: header_footer(c, d, case)
    doc.onLaterPages = lambda c, d: header_footer(c, d, case)

    story = []
    w = 170*mm

    # ── Cover page ──
    story.append(Spacer(1, 8*mm))
    cover_data = [[
        Table([
            [Paragraph("SOC INCIDENT REPORT", styles["CoverTitle"])],
            [Spacer(1, 3*mm)],
            [Paragraph(case.get("title", "Untitled Incident"), styles["CoverSub"])],
            [Spacer(1, 6*mm)],
            [Table([
                [Paragraph(f"Case Number", styles["CoverMeta"]), Paragraph(case.get("case_number", "N/A"), styles["CoverMeta"])],
                [Paragraph(f"Severity", styles["CoverMeta"]), Paragraph(case.get("severity", "N/A"), styles["CoverMeta"])],
                [Paragraph(f"Incident Type", styles["CoverMeta"]), Paragraph(case.get("incident_type", "N/A"), styles["CoverMeta"])],
                [Paragraph(f"Status", styles["CoverMeta"]), Paragraph(case.get("status", "Open"), styles["CoverMeta"])],
                [Paragraph(f"Analyst", styles["CoverMeta"]), Paragraph(case.get("analyst_name", "N/A"), styles["CoverMeta"])],
                [Paragraph(f"Customer", styles["CoverMeta"]), Paragraph(case.get("customer_name", "N/A"), styles["CoverMeta"])],
                [Paragraph(f"Classification", styles["CoverMeta"]), Paragraph(case.get("classification", "TLP:AMBER"), styles["CoverMeta"])],
                [Paragraph(f"Date", styles["CoverMeta"]), Paragraph(datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"), styles["CoverMeta"])],
            ], colWidths=[40*mm, 80*mm])],
        ], colWidths=[w])
    ]]
    cover_tbl = Table(cover_data, colWidths=[w])
    cover_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), NAVY),
        ("TOPPADDING", (0, 0), (-1, -1), 20),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 20),
        ("LEFTPADDING", (0, 0), (-1, -1), 16),
        ("RIGHTPADDING", (0, 0), (-1, -1), 16),
    ]))
    story.append(cover_tbl)
    story.append(PageBreak())

    # ── Executive Summary ──
    story.append(section_banner("Executive Summary", styles))
    story.append(Spacer(1, 3*mm))
    exec_text = case.get("ai_executive_summary", "AI summary not generated yet.")
    for para in exec_text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), styles["BodyText"]))
            story.append(Spacer(1, 2*mm))

    # Severity score block
    score = case.get("ai_severity_score", 0)
    score_color = RED_CRIT if score >= 75 else ORANGE_HI if score >= 50 else YELLOW_MD if score >= 25 else GREEN_LO
    story.append(Spacer(1, 4*mm))
    score_data = [[
        Paragraph("AI RISK SCORE", ParagraphStyle("ScoreLabel", fontSize=8, textColor=colors.HexColor("#90CAF9"), fontName="Helvetica-Bold")),
        Paragraph(f"{score}/100", ParagraphStyle("ScoreNum", fontSize=20, textColor=score_color, fontName="Helvetica-Bold")),
        Paragraph(case.get("ai_severity_reasoning", "")[:200], ParagraphStyle("ScoreText", fontSize=9, textColor=colors.HexColor("#CFD8DC"), fontName="Helvetica")),
    ]]
    score_tbl = Table(score_data, colWidths=[28*mm, 25*mm, 117*mm])
    score_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), SLATE),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(score_tbl)

    # ── IOC Table ──
    story.append(Spacer(1, 6*mm))
    story.append(section_banner("Indicators of Compromise (IOCs)", styles))
    story.append(Spacer(1, 3*mm))
    if iocs:
        story.append(ioc_table(iocs, vt_results, styles))
    else:
        story.append(Paragraph("No IOCs recorded for this case.", styles["BodyText"]))

    # ── MITRE ATT&CK ──
    story.append(Spacer(1, 6*mm))
    story.append(section_banner("MITRE ATT&CK Mapping", styles))
    story.append(Spacer(1, 3*mm))
    story.append(mitre_table(mitre, styles))

    # ── Timeline ──
    story.append(Spacer(1, 6*mm))
    story.append(section_banner("Attack Timeline", styles))
    story.append(Spacer(1, 3*mm))
    story.append(timeline_table(timeline, styles))

    # ── Technical Analysis ──
    story.append(Spacer(1, 6*mm))
    story.append(section_banner("Technical Analysis", styles))
    story.append(Spacer(1, 3*mm))
    tech_text = case.get("ai_technical_summary", "AI technical summary not generated yet.")
    for para in tech_text.split("\n\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), styles["BodyText"]))
            story.append(Spacer(1, 2*mm))

    # ── Commands Run ──
    if case.get("commands_run"):
        story.append(Spacer(1, 4*mm))
        story.append(section_banner("Commands Run During Investigation", styles))
        story.append(Spacer(1, 3*mm))
        for cmd in case.get("commands_run", "").split("\n"):
            if cmd.strip():
                story.append(Paragraph(cmd.strip(), styles["Mono"]))
                story.append(Spacer(1, 1*mm))

    # ── Recommendations ──
    story.append(Spacer(1, 6*mm))
    story.append(section_banner("Recommendations", styles))
    story.append(Spacer(1, 3*mm))
    recs = case.get("recommendations", case.get("ai_technical_summary", ""))
    for line in recs.split("\n"):
        if line.strip():
            story.append(Paragraph(line.strip(), styles["BodyText"]))
            story.append(Spacer(1, 1.5*mm))

    doc.build(story)
    return buf.getvalue()

# ─── DOCX Generator ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def add_heading(doc, text, level, color_hex="0D1B2A"):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.bold = True
    return p

def generate_docx(case: Dict) -> bytes:
    doc = DocxDocument()
    iocs = _parse_json_field(case.get("iocs", "[]"), [])
    timeline = _parse_json_field(case.get("timeline_events", "[]"), [])
    mitre = _parse_json_field(case.get("mitre_techniques", "[]"), [])
    vt_results = _parse_json_field(case.get("vt_results", "{}"), {})

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOC INCIDENT REPORT")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(case.get("title", "Untitled Incident"))
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x00, 0xB4, 0xD8)

    doc.add_paragraph()

    # Meta table
    meta_data = [
        ("Case Number", case.get("case_number", "N/A")),
        ("Severity", case.get("severity", "N/A")),
        ("Incident Type", case.get("incident_type", "N/A")),
        ("Status", case.get("status", "Open")),
        ("Analyst", case.get("analyst_name", "N/A")),
        ("Customer", case.get("customer_name", "N/A")),
        ("Classification", case.get("classification", "TLP:AMBER")),
        ("Date", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
    ]
    tbl = doc.add_table(rows=len(meta_data), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(meta_data):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_bg(tbl.rows[i].cells[0], "1E2D3D")
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_page_break()

    # Sections
    sections = [
        ("Executive Summary", case.get("ai_executive_summary", "AI summary not generated.")),
        ("Technical Analysis", case.get("ai_technical_summary", "AI technical summary not generated.")),
        ("Recommendations", case.get("recommendations", "")),
    ]
    for heading, content in sections:
        h = doc.add_heading(heading, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        for para in content.split("\n\n"):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(10)

    # IOC Table
    h = doc.add_heading("Indicators of Compromise (IOCs)", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
    if iocs:
        tbl = doc.add_table(rows=1 + len(iocs), cols=4)
        tbl.style = "Table Grid"
        headers = ["IOC", "Type", "VT Verdict", "Detections"]
        for j, hdr in enumerate(headers):
            tbl.rows[0].cells[j].text = hdr
            tbl.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
            set_cell_bg(tbl.rows[0].cells[j], "1E2D3D")
            tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, ioc in enumerate(iocs):
            vt = vt_results.get(ioc, {})
            verdict = vt.get("verdict", "not_checked")
            det = f"{vt.get('malicious_count', '-')}/{vt.get('total_engines', '-')}" if vt.get("total_engines") else "—"
            row_data = [ioc[:50], vt.get("type", "—"), verdict.upper(), det]
            for j, val in enumerate(row_data):
                tbl.rows[i + 1].cells[j].text = val
    else:
        doc.add_paragraph("No IOCs recorded.")

    # MITRE Table
    h = doc.add_heading("MITRE ATT&CK Mapping", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
    if mitre:
        tbl = doc.add_table(rows=1 + len(mitre), cols=4)
        tbl.style = "Table Grid"
        headers = ["Technique ID", "Name", "Tactic", "Evidence"]
        for j, hdr in enumerate(headers):
            tbl.rows[0].cells[j].text = hdr
            tbl.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
            set_cell_bg(tbl.rows[0].cells[j], "1E2D3D")
            tbl.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for i, tech in enumerate(mitre):
            row_data = [tech.get("technique_id", ""), tech.get("technique_name", ""), tech.get("tactic", ""), tech.get("evidence", "")[:100]]
            for j, val in enumerate(row_data):
                tbl.rows[i + 1].cells[j].text = val

    # Timeline
    if timeline:
        h = doc.add_heading("Attack Timeline", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        for ev in timeline:
            p = doc.add_paragraph(str(ev), style="List Bullet")

    # Commands
    if case.get("commands_run"):
        h = doc.add_heading("Commands Run", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        p = doc.add_paragraph()
        run = p.add_run(case.get("commands_run", ""))
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
